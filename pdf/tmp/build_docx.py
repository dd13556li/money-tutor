# -*- coding: utf-8 -*-
"""
金錢小達人 — 教育部第八屆特殊教育教材教具競賽 完整申請文件產生器
申請組別：身心障礙類 國民中學及高級中等學校組
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── 工具函數 ────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def bold_run(para, text, size=12, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'DFKai-SB'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def normal_run(para, text, size=12, bold=False):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'DFKai-SB'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
    return run

def add_heading(doc, text, level=1, size=14, color=(0,51,102)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'DFKai-SB'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
    run.font.color.rgb = RGBColor(*color)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    return p

def add_para(doc, text, size=12, bold=False, indent=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'DFKai-SB'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    return p

def cell_text(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'DFKai-SB'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(18)

def cell_multiline(cell, lines, size=11, header_bold=False):
    cell.text = ''
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.bold = (i == 0 and header_bold)
        run.font.size = Pt(size)
        run.font.name = 'DFKai-SB'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(18)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type='page')

# ─── 文件初始化 ─────────────────────────────────────────────

from docx.enum.text import WD_BREAK
from docx.oxml.ns import nsmap

doc = Document()

# 頁面設定：A4，邊界
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(3.17)
section.bottom_margin = Cm(3.17)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)

# 預設段落樣式
style = doc.styles['Normal']
style.font.name = 'DFKai-SB'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
style.paragraph_format.line_spacing = Pt(22)

# ═══════════════════════════════════════════════════════
# 封面頁
# ═══════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('第八屆教育部獎助研發特殊教育教材教具競賽')
run.bold = True; run.font.size = Pt(16)
run.font.name = 'DFKai-SB'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
run.font.color.rgb = RGBColor(0,51,102)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('金錢小達人')
run.bold = True; run.font.size = Pt(28)
run.font.name = 'DFKai-SB'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
run.font.color.rgb = RGBColor(0,102,51)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('互動式特殊教育生活數學教學軟體')
run.bold = True; run.font.size = Pt(18)
run.font.name = 'DFKai-SB'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DFKai-SB')
run.font.color.rgb = RGBColor(51,51,51)

for _ in range(2):
    doc.add_paragraph()

# 基本資訊框
tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ('申請組別', '■ 身心障礙類：■ 國民中學及高級中等學校組'),
    ('適用對象', '國民中學至高級中等學校特殊需求學生'),
    ('作品類型', '科技輔助教學軟體（網頁應用程式）'),
    ('申請日期', '中華民國 115 年 3 月'),
]
for i, (k, v) in enumerate(info):
    cell_text(tbl.rows[i].cells[0], k, size=12, bold=True)
    cell_text(tbl.rows[i].cells[1], v, size=12)
    set_cell_bg(tbl.rows[i].cells[0], 'D9E1F2')
    tbl.rows[i].cells[0].width = Cm(4.5)
    tbl.rows[i].cells[1].width = Cm(11.0)

# 分頁
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ═══════════════════════════════════════════════════════
# 附表一 ─ 申請表
# ═══════════════════════════════════════════════════════

add_heading(doc, '附表一　教育部獎助研發特殊教育教材教具競賽申請表', 1, 14)
add_para(doc, '作品編號：　　　　　　　　（由承辦單位編列）', size=11)
doc.add_paragraph()

t1 = doc.add_table(rows=16, cols=4)
t1.style = 'Table Grid'

def merge_row_cells(table, row_idx, start_col, end_col):
    cells = table.rows[row_idx].cells
    merged = cells[start_col].merge(cells[end_col])
    return merged

# 列0：組別（跨欄）
r0 = t1.rows[0]
r0.cells[0].merge(r0.cells[1])
cell_text(r0.cells[0], '參賽組別', bold=True, size=11)
r0.cells[2].merge(r0.cells[3])
cell_multiline(r0.cells[2], [
    '■ 身心障礙類：□ 學前教育組　□ 國民小學組',
    '                    ■ 國民中學及高級中等學校組',
    '□ 資賦優異類：□ 國民小學組　□ 國民中學及高級中等學校組',
], size=11)
set_cell_bg(r0.cells[0], 'D9E1F2')

# 列1：作品名稱
r1 = t1.rows[1]
r1.cells[0].merge(r1.cells[1])
cell_text(r1.cells[0], '作品名稱', bold=True, size=11)
r1.cells[2].merge(r1.cells[3])
cell_text(r1.cells[2], '金錢小達人——互動式特殊教育生活數學教學軟體', size=11)
set_cell_bg(r1.cells[0], 'D9E1F2')

# 列2：適用對象
r2 = t1.rows[2]
r2.cells[0].merge(r2.cells[1])
cell_text(r2.cells[0], '適用對象\n（教育階段/類別）', bold=True, size=11)
r2.cells[2].merge(r2.cells[3])
cell_text(r2.cells[2], '國民中學及高級中等學校階段 / 身心障礙類（含輕中度智能障礙、自閉症、學習障礙）', size=11)
set_cell_bg(r2.cells[0], 'D9E1F2')

# 列3：申請人標頭
r3 = t1.rows[3]
cell_text(r3.cells[0], '姓名', bold=True, size=11)
cell_text(r3.cells[1], '', size=11)
cell_text(r3.cells[2], '身分證明文件號碼', bold=True, size=11)
cell_text(r3.cells[3], '', size=11)
set_cell_bg(r3.cells[0], 'D9E1F2')
set_cell_bg(r3.cells[2], 'D9E1F2')

# 列4：服務單位
r4 = t1.rows[4]
cell_text(r4.cells[0], '服務單位', bold=True, size=11)
cell_text(r4.cells[1], '', size=11)
cell_text(r4.cells[2], '職稱', bold=True, size=11)
cell_text(r4.cells[3], '', size=11)
set_cell_bg(r4.cells[0], 'D9E1F2')
set_cell_bg(r4.cells[2], 'D9E1F2')

# 列5：照片提示（跨列）
r5 = t1.rows[5]
r5.cells[0].merge(r5.cells[1])
cell_multiline(r5.cells[0], [
    '申請獎助備註說明',
    '',
    '一、是否曾獲其他政府機關(構)、學校競賽獎助',
    '    □否    □是，獎助名稱：',
    '',
    '二、是否已另向其他政府機關(構)申請獎助（若否，免填第三項）',
    '    □否    □是，獎助名稱：',
    '',
    '三、若同時獲得本項及其他全國性獎助，請勾選意願：',
    '    □ 願主動聲明放棄其他獎助',
    '    □ 願主動聲明放棄本項獎助',
    '',
    '四、是否為二人以上共同作品：',
    '    □ 否，請檢具附表二著作權聲明及授權書',
    '    □ 是，請檢具附表二（每人各一份）及附表四共同作者同意書',
], size=11)
set_cell_bg(r5.cells[0], 'FAFAFA')
r5.cells[2].merge(r5.cells[3])
cell_multiline(r5.cells[2], [
    '請黏貼最近一年內',
    '一吋脫帽半身照片',
    '（或彩色列印）',
    '',
    '',
    '',
    '',
    '',
], size=11)

# 列6：聯絡資訊
r6 = t1.rows[6]
r6.cells[0].merge(r6.cells[3])
cell_multiline(r6.cells[0], [
    '本人均已熟知並切結遵行本次申請內容及相關規定。',
    '',
    '申請人簽名：　　　　　　　　　　　　',
    '聯絡地址：',
    '聯絡電話：',
    '電子郵件信箱：',
], size=11)

# 列7-15：空白備用
for i in range(7, 16):
    t1.rows[i].cells[0].merge(t1.rows[i].cells[3])
    cell_text(t1.rows[i].cells[0], '（備用空白欄）' if i == 15 else '', size=10)

# 修正列高
set_row_height(t1.rows[5], 7.0)
set_row_height(t1.rows[6], 3.5)

# 刪除多餘行
for i in range(15, 6, -1):
    if i > 6:
        tr = t1.rows[i]._tr
        tr.getparent().remove(tr)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ═══════════════════════════════════════════════════════
# 附表三 ─ 創作說明
# ═══════════════════════════════════════════════════════

add_heading(doc, '附表三　創作說明', 1, 14)

t3 = doc.add_table(rows=11, cols=2)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

HEADER_W = Cm(3.0)
CONTENT_W = Cm(12.5)

def t3_row(table, row_idx, header, content_lines):
    cell_text(table.rows[row_idx].cells[0], header, bold=True, size=11)
    set_cell_bg(table.rows[row_idx].cells[0], 'D9E1F2')
    cell_multiline(table.rows[row_idx].cells[1], content_lines, size=11)
    table.rows[row_idx].cells[0].width = HEADER_W
    table.rows[row_idx].cells[1].width = CONTENT_W

t3_row(t3, 0, '參賽組別', [
    '■ 身心障礙類：■ 國民中學及高級中等學校組'
])

t3_row(t3, 1, '作品名稱', [
    '金錢小達人——互動式特殊教育生活數學教學軟體'
])

t3_row(t3, 2, '適用對象\n（教育階段/類別）', [
    '國民中學及高級中等學校階段身心障礙學生',
    '包含：輕度至中度智能障礙、自閉症、學習障礙、多重障礙等特殊需求學生',
    '（部分單元經調整後亦適用國小中高年級特殊需求學生）',
])

t3_row(t3, 3, '作品用途', [
    '本軟體為一套以網頁為基礎的互動式特殊教育生活數學電腦輔助教學軟體，',
    '適用於特教教師課堂教學、資源班個別化教學及家長家庭練習。',
    '透過 18 個情境單元，系統性培養學生貨幣辨識、金額計算、',
    '找零付款及日常消費等功能性生活技能，',
    '並提供對應的印刷式作業單，支援紙本評量與課後複習。',
])

t3_row(t3, 4, '設計動機', [
    '一、現場觀察：',
    '    特殊需求學生在「生活數學」領域，特別是金錢使用技能上，往往因',
    '    抽象概念難以內化、缺乏安全練習環境而產生學習困境。研究者長期',
    '    觀察教學現場，發現學生雖然在課堂上學過面額，卻在ATM操作、',
    '    自動販賣機購物、搭車購票等真實場景中仍感到無助與焦慮。',
    '',
    '二、課程需求：',
    '    依據「十二年國教特殊教育課程實施規範」，國中特殊教育應強調',
    '    「生活技能」與「職業準備」，高中則進一步連結「社區參與」與',
    '    「轉銜教育」。坊間缺乏對應課綱且適用於中學階段的系統性金錢',
    '    技能電腦輔助教材。',
    '',
    '三、通用設計理念（UDL）：',
    '    本軟體依據課程通用設計原則（UDL）三大核心——多元表徵、多元',
    '    參與、多元行動與表達——設計三種難度、語音引導、視覺提示及',
    '    輔助點擊模式，使不同能力程度的學生均能獲得適切的學習支持。',
])

t3_row(t3, 5, '內容概述\n（含製作過程）', [
    '【軟體架構】',
    '本軟體以 HTML5 + CSS3 + JavaScript 開發，採無框架純前端架構，',
    '無需安裝，開啟瀏覽器即可使用，支援電腦、平板與手機多裝置。',
    '開發總計 18 個主程式（各約 3,000～15,000 行），',
    '程式碼總量超過 15 萬行，歷時約 4 個月完成。',
    '',
    '【三大系列 18 個教學單元】',
    '',
    '▍ F 系列：數學基礎（6 個單元）',
    '  F1 一對一對應：拖曳圖示配對，建立數量概念基礎。',
    '  F2 唱數：逐步數數，含口語語音同步引導，三種難度。',
    '  F3 數字認讀：看圖選數字，含自訂主題圖片上傳功能。',
    '  F4 數字排序：拖曳數字完成由小到大或由大到小排列。',
    '  F5 量比較：比較兩組物品多寡，建立大小比較概念。',
    '  F6 數的組成：以合成與分解概念理解數的結構。',
    '',
    '▍ C 系列：貨幣認知（6 個單元）',
    '  C1 認識錢幣：辨識台灣 9 種面額（硬幣4種、鈔票5種），',
    '     三種難度（看圖選字 / 看字選幣 / 聽音選幣）。',
    '  C2 數錢：點擊螢幕上的錢幣逐枚數算，語音同步播報。',
    '  C3 換錢：以不同面額組合湊出等值金額，建立兌換概念。',
    '  C4 付款：從錢包中拖曳正確金額至付款區完成付款。',
    '  C5 夠不夠：判斷現有金額能否購買指定商品。',
    '  C6 找零：計算應找回的零錢，並正確放置至找零區。',
    '',
    '▍ A 系列：生活應用（6 個單元）',
    '  A1 自動販賣機：完整模擬選飲料→投幣→確認→取物→取找零。',
    '     特色：「先投幣模式」模擬真實販賣機流程（投幣後燈號亮起）。',
    '  A2 理髮廳自助機：模擬服務付款與取號等候流程。',
    '  A3 麥當勞點餐機：支援70種餐點、四類別、客製自訂餐點上傳。',
    '  A4 超市購物：12種商店類型×119種商品，含完整結帳找零。',
    '  A5 ATM 操作：完整模擬提款、存款、查詢、轉帳四種交易，',
    '     含 PIN 密碼輸入、交易步驟、明細列印、退卡全流程。',
    '  A6 火車售票機：模擬選擇車站、票種、張數，完成購票付款。',
    '',
    '【三種難度設計】',
    '  簡單模式：全程語音引導 + 黃色光暈提示，錯誤自動引導。',
    '  普通模式：半引導，錯誤 3 次後顯示提示。',
    '  困難模式：完全自主操作，無自動提示。',
    '',
    '【輔助點擊模式（無障礙設計）】',
    '  啟用後，畫面任何位置點擊即執行下一步操作，',
    '  系統自動建立操作佇列，配合語音引導，',
    '  適用於動作協調困難或認知負荷較高的學生。',
    '',
    '【作業單系統】',
    '  18 個單元各有對應印刷式作業單，共 10 種以上題型，',
    '  含「看圖填空」、「選出正確金額」、「計算找零」等，',
    '  可由教師調整難度、題數與題型後一鍵列印。',
    '',
    '【獎勵系統】',
    '  教師可為學生累積點數、上傳照片、設定獎項，',
    '  增強學習動機，適合個別化教育計畫（IEP）搭配使用。',
])

t3_row(t3, 6, '使用說明\n及效果', [
    '【使用方式】',
    '1. 開啟電腦/平板瀏覽器，進入軟體首頁，選擇單元，即可開始教學。',
    '2. 教師於設定頁選擇難度（簡單/普通/困難）、題數及任務類型。',
    '3. 可開啟「輔助點擊模式」降低操作門檻，適用於需要高支持學生。',
    '4. 每題完成後播放語音回饋，全部完成後顯示成績統計與煙火動畫。',
    '5. 作業單：於各單元設定頁點擊「作業單」按鈕，設定後列印即可。',
    '',
    '【課程融入建議】',
    '  本軟體可融入：生活數學、特殊需求領域（生活管理）、',
    '  轉銜教育、職業教育、社區參與訓練等課程。',
    '  建議搭配實體錢幣操作、社區購物體驗，形成「數位練習→',
    '  紙本評量→實地應用」的完整學習歷程。',
    '',
    '【預期教學成效】',
    '1. 提升學生對台灣幣值的辨識與計算能力。',
    '2. 建立在自動販賣機、ATM、商店等真實場景的操作信心。',
    '3. 透過反覆練習降低學生面對消費情境的焦慮。',
    '4. 提供教師多樣化的評量工具，節省自製教材時間。',
    '5. 支援個別化教育計畫（IEP）之功能性技能目標設定。',
])

t3_row(t3, 7, '附錄', [
    '附錄一：特殊教育教學活動設計（附表三之一）',
    '附錄二：軟體各單元操作截圖（略，另附光碟）',
    '附錄三：作業單範例（略，另附光碟）',
    '附錄四：學生試用回饋記錄（略，另附光碟）',
])

t3_row(t3, 8, '製作費用', [
    '軟體開發費：新臺幣 0 元（研究者自主開發）',
    '影像素材費：新臺幣 0 元（原創繪製或合法授權免費素材）',
    '音效費用：新臺幣 0 元（合法授權免費音效）',
    '硬體設備：由學校現有電腦/平板設備支應，無額外費用',
])

t3_row(t3, 9, '經費來源', ['自籌'])

t3_row(t3, 10, '備註', [
    '本軟體為網頁應用程式，無需安裝，可直接於現代瀏覽器執行。',
    '建議使用 Chrome 或 Edge 瀏覽器以獲得最佳效果。',
    '軟體已完整測試於 Windows 電腦、iPad 及 Android 平板。',
])

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ═══════════════════════════════════════════════════════
# 附表三之一 ─ 教學活動設計（單元一）
# ═══════════════════════════════════════════════════════

add_heading(doc, '附表三之一　特殊教育教學活動設計', 1, 14)
add_heading(doc, '教學設計（一）：貨幣辨識與日常消費技能', 2, 13, (0,102,51))
add_para(doc, '——以自動販賣機與超市購物為情境（A1、A4 單元）', size=12)
doc.add_paragraph()

# 基本資訊表
ti = doc.add_table(rows=14, cols=4)
ti.style = 'Table Grid'

def ti_row(row_idx, labels_values):
    row = ti.rows[row_idx]
    for col, (label, value) in enumerate(labels_values):
        is_label = (col % 2 == 0)
        cell_text(row.cells[col], label if is_label else value,
                  bold=is_label, size=11)
        if is_label:
            set_cell_bg(row.cells[col], 'D9E1F2')
            row.cells[col].width = Cm(2.8)
        else:
            row.cells[col].width = Cm(5.2)

ti_row(0, [('領域/科目', ''), ('', '')])
ti.rows[0].cells[0].merge(ti.rows[0].cells[1])
ti.rows[0].cells[2].merge(ti.rows[0].cells[3])
cell_text(ti.rows[0].cells[0], '領域/科目', bold=True, size=11)
set_cell_bg(ti.rows[0].cells[0], 'D9E1F2')
cell_multiline(ti.rows[0].cells[2], [
    '■ 數學領域　■ 融入特殊需求領域（生活管理）'
])

ti_row(1, [('形式', ''), ('', '')])
ti.rows[1].cells[0].merge(ti.rows[1].cells[1])
ti.rows[1].cells[2].merge(ti.rows[1].cells[3])
cell_text(ti.rows[1].cells[0], '形式', bold=True, size=11)
set_cell_bg(ti.rows[1].cells[0], 'D9E1F2')
cell_text(ti.rows[1].cells[2], '■ 融入　生活技能/職業教育　領域/科目')

ti_row(2, [('實施型態', ''), ('', '')])
ti.rows[2].cells[0].merge(ti.rows[2].cells[1])
ti.rows[2].cells[2].merge(ti.rows[2].cells[3])
cell_text(ti.rows[2].cells[0], '實施型態', bold=True, size=11)
set_cell_bg(ti.rows[2].cells[0], 'D9E1F2')
cell_multiline(ti.rows[2].cells[2], [
    '■ 身心障礙分散式資源班　■ 集中式特教班',
    '□ 身心障礙巡迴輔導班',
])

ti_row(3, [('單元名稱', ''), ('版本', '')])
ti.rows[3].cells[0].merge(ti.rows[3].cells[0])
cell_text(ti.rows[3].cells[0], '單元名稱', bold=True, size=11)
set_cell_bg(ti.rows[3].cells[0], 'D9E1F2')
cell_text(ti.rows[3].cells[1], '從購物開始——辨幣、付款、找零實戰演練')
cell_text(ti.rows[3].cells[2], '版本', bold=True, size=11)
set_cell_bg(ti.rows[3].cells[2], 'D9E1F2')
cell_text(ti.rows[3].cells[3], '自編（金錢小達人 C1~C6、A1、A4 單元）')

ti_row(4, [('學習階段/年級', ''), ('教學設計者', '')])
cell_text(ti.rows[4].cells[0], '學習階段/年級', bold=True, size=11)
set_cell_bg(ti.rows[4].cells[0], 'D9E1F2')
cell_text(ti.rows[4].cells[1], '第四至第五階段 / 七至十二年級')
cell_text(ti.rows[4].cells[2], '教學設計者', bold=True, size=11)
set_cell_bg(ti.rows[4].cells[2], 'D9E1F2')
cell_text(ti.rows[4].cells[3], '特教教師（請填入姓名）')

ti_row(5, [('學習功能', ''), ('教學時間', '')])
cell_text(ti.rows[5].cells[0], '學習功能', bold=True, size=11)
set_cell_bg(ti.rows[5].cells[0], 'D9E1F2')
cell_multiline(ti.rows[5].cells[1], ['■ 輕微缺損　■ 嚴重缺損'])
cell_text(ti.rows[5].cells[2], '教學時間', bold=True, size=11)
set_cell_bg(ti.rows[5].cells[2], 'D9E1F2')
cell_text(ti.rows[5].cells[3], '共 4 節 / 本設計為第 1-4 節（每節45分鐘）')

ti_row(6, [('教學地點', ''), ('', '')])
ti.rows[6].cells[0].merge(ti.rows[6].cells[1])
ti.rows[6].cells[2].merge(ti.rows[6].cells[3])
cell_text(ti.rows[6].cells[0], '教學地點', bold=True, size=11)
set_cell_bg(ti.rows[6].cells[0], 'D9E1F2')
cell_text(ti.rows[6].cells[2], '電腦教室、資源班教室或具備平板設備之教室')

ti_row(7, [('特殊需求', ''), ('', '')])
ti.rows[7].cells[0].merge(ti.rows[7].cells[1])
ti.rows[7].cells[2].merge(ti.rows[7].cells[3])
cell_text(ti.rows[7].cells[0], '特殊需求', bold=True, size=11)
set_cell_bg(ti.rows[7].cells[0], 'D9E1F2')
cell_text(ti.rows[7].cells[2], '■ 生活管理　■ 輔助科技應用　■ 職業教育')

ti_row(8, [('議題融入', ''), ('', '')])
ti.rows[8].cells[0].merge(ti.rows[8].cells[1])
ti.rows[8].cells[2].merge(ti.rows[8].cells[3])
cell_text(ti.rows[8].cells[0], '議題融入', bold=True, size=11)
set_cell_bg(ti.rows[8].cells[0], 'D9E1F2')
cell_text(ti.rows[8].cells[2], '■ 科技教育　■ 資訊教育　■ 生涯規劃教育　■ 家庭教育')

# 設計理念
ti_row(9, [('設計理念', ''), ('', '')])
ti.rows[9].cells[0].merge(ti.rows[9].cells[1])
ti.rows[9].cells[2].merge(ti.rows[9].cells[3])
cell_text(ti.rows[9].cells[0], '設計理念', bold=True, size=11)
set_cell_bg(ti.rows[9].cells[0], 'D9E1F2')
cell_multiline(ti.rows[9].cells[2], [
    '國中至高中特殊需求學生面臨的重要任務是「功能性生活技能」的建立，',
    '尤其是金錢管理技能對其日常生活獨立、社區參與及未來就業均至關重要。',
    '本設計以「認知建立→情境模擬→數位練習→實地應用」為學習歷程架構，',
    '結合金錢小達人教學軟體，透過多感官互動（視覺動畫、聽覺語音、觸覺操作），',
    '讓學生在安全的數位環境中反覆練習「辨識錢幣→數算金額→付款找零」',
    '的完整消費流程，並配合作業單進行紙本評量，達成功能性技能學習目標。',
    '本教學採多層次教學策略，同一教室中不同程度學生可同步進行不同難度練習，',
    '完全符合融合教育與個別化教育計畫（IEP）精神。',
], size=11)
set_row_height(ti.rows[9], 4.5)

# 刪除多餘列
for i in range(13, 9, -1):
    tr = ti.rows[i]._tr
    tr.getparent().remove(tr)

doc.add_paragraph()

# ── 學生能力描述
add_heading(doc, '學生能力描述', 2, 12, (0,51,102))

ts = doc.add_table(rows=5, cols=5)
ts.style = 'Table Grid'
ts_headers = ['學生', '特教類別', '一般現況能力', '領域能力表現（數學/生活管理）', '課程調整']
for i, h in enumerate(ts_headers):
    cell_text(ts.rows[0].cells[i], h, bold=True, size=10)
    set_cell_bg(ts.rows[0].cells[i], 'D9E1F2')

students = [
    ('甲生', '輕度智能障礙',
     '識字約國小四年級程度；能數算到100；具基本自理能力；對電腦操作有高度興趣',
     '能辨識1、5、10元硬幣；知道面額但計算組合時常出錯；對付款場景感到焦慮',
     '簡化：軟體設定簡單模式+輔助點擊；限定硬幣種類；作業單使用圈選式'),
    ('乙生', '自閉症（輕度）',
     '認知能力接近同齡；語言理解良好但表達較弱；對固定流程有強烈依賴',
     '能計算金額，但遇到不熟悉的情境（如新店家）時常出現拒絕行為',
     '減量+歷程調整：從熟悉情境（自販機）開始，逐步擴展至新場景；軟體普通模式'),
    ('丙生', '學習障礙',
     '語文能力正常；數學計算有困難；學習動機良好；能獨立使用電腦',
     '知道面額意義，但進行多位數加減計算時常出錯；找零時需要較長思考時間',
     '歷程調整：使用軟體提示線索功能；提供計算輔具（計算機）；普通或困難模式'),
    ('丁生', '中度智能障礙',
     '識字有限，以圖示溝通為主；動作發展接近同齡；語言表達以單詞為主',
     '能辨識一元硬幣，尚無法辨識其他面額；對購物情境有參與意願',
     '替代+減量：軟體簡單模式+輔助點擊；僅練習辨識與選擇，不要求計算'),
]
for r, (s, cat, gen, domain, adj) in enumerate(students):
    row = ts.rows[r+1]
    for c, text in enumerate([s, cat, gen, domain, adj]):
        cell_text(row.cells[c], text, size=10)

doc.add_paragraph()

# ── 設計依據
add_heading(doc, '設計依據', 2, 12, (0,51,102))
td = doc.add_table(rows=5, cols=2)
td.style = 'Table Grid'
design_items = [
    ('核心素養\n具體內涵', [
        '數-J-A2 具備基本算術操作能力，理解數量關係，在生活情境中運用。',
        '數-U-A2 能運用數學解決日常生活問題，並與社會情境連結。',
        '特管-B2 應用適當技巧解決日常生活問題（生活管理領域）。',
    ]),
    ('學科領域\n學習重點', [
        '學習表現：n-IV-5 能處理日常生活中的金錢計算問題。',
        '        　n-IV-6 能運用估算解決日常生活問題。',
        '學習內容：N-7-1 生活中的整數（含金錢使用）。',
        '調整後表現：n-IV-5-1 能辨識硬幣面額並計算組合金額。',
        '           n-IV-5-2 能在引導下完成付款與確認找零。',
    ]),
    ('特殊需求\n領域', [
        '學習重點：特管3-IV-1 能使用金錢完成日常消費活動（生活管理）。',
        '         特管3-IV-2 能處理日常購物的付款與找零。',
        '         特輔-B2 能應用輔助科技進行學習與生活管理。',
    ]),
    ('議題學習\n主題', [
        '生涯規劃教育：探索日常消費技能對未來獨立生活的重要性。',
        '科技教育：運用數位科技工具輔助學習生活數學技能。',
        '家庭教育：連結家庭日常消費情境，促進家庭類化練習。',
    ]),
    ('議題實質\n內涵', [
        '生J5 了解並嘗試進行個人的金錢管理。',
        '科J4 能運用數位工具解決生活中的學習問題。',
        '家J6 認識家庭日常消費與財務管理的基本概念。',
    ]),
]
for i, (hdr, lines) in enumerate(design_items):
    cell_text(td.rows[i].cells[0], hdr, bold=True, size=11)
    set_cell_bg(td.rows[i].cells[0], 'D9E1F2')
    cell_multiline(td.rows[i].cells[1], lines, size=11)
    td.rows[i].cells[0].width = Cm(3.0)
    td.rows[i].cells[1].width = Cm(12.5)

doc.add_paragraph()

# ── 學習調整
add_heading(doc, '學習調整', 2, 12, (0,51,102))
tadj = doc.add_table(rows=4, cols=3)
tadj.style = 'Table Grid'
cell_text(tadj.rows[0].cells[0], '調整方式', bold=True, size=11)
cell_text(tadj.rows[0].cells[1], '調整策略說明', bold=True, size=11)
cell_text(tadj.rows[0].cells[2], '對應勾選', bold=True, size=11)
for c in range(3): set_cell_bg(tadj.rows[0].cells[c], 'D9E1F2')

adj_data = [
    ('內容', [
        '・針對嚴重缺損學生：簡化學習目標，聚焦於辨識常用面額（1、5、10元）',
        '・針對輕微缺損學生：減量計算範圍（以整數、50元以內為主）',
        '・替代：以圖示操作取代抽象數字計算（甲生、丁生）',
    ], '■ 簡化　■ 減量\n■ 分解　■ 替代'),
    ('歷程', [
        '・工作分析：將購物流程分解為可獨立完成的步驟（辨幣→數算→選取→付款→找零）',
        '・多元感官：軟體提供視覺動畫、聽覺語音、觸覺操作三感官同步支持',
        '・提示線索：簡單模式黃色光暈高亮提示，普通模式錯誤三次出現提示',
        '・多層次教學：同一情境中不同學生操作不同難度（甲丁簡單/乙普通/丙困難）',
    ], '■ 提示線索　■ 多元感官\n■ 工作分析　■ 多層次教學'),
    ('評量', [
        '・作答方式調整：圈選式題目（甲生、丁生）；計算題提供計算機輔具（丙生）',
        '・試題調整：簡化金額範圍；選項數量減少（甲生：3選項→2選項）',
        '・評量工具：軟體即時回饋、紙本作業單、口頭問答、實作評量多元並用',
    ], '■ 試題調整　■ 作答方式調整\n■ 輔具：計算機、平板'),
]
for i, (way, lines, check) in enumerate(adj_data):
    cell_text(tadj.rows[i+1].cells[0], way, bold=True, size=11)
    set_cell_bg(tadj.rows[i+1].cells[0], 'FAFAFA')
    cell_multiline(tadj.rows[i+1].cells[1], lines, size=11)
    cell_text(tadj.rows[i+1].cells[2], check, size=11)
    tadj.rows[i+1].cells[0].width = Cm(1.5)
    tadj.rows[i+1].cells[1].width = Cm(10.5)
    tadj.rows[i+1].cells[2].width = Cm(3.5)

doc.add_paragraph()

# ── 學習目標
add_heading(doc, '學習目標', 2, 12, (0,51,102))
tobj = doc.add_table(rows=8, cols=3)
tobj.style = 'Table Grid'
cell_text(tobj.rows[0].cells[0], '代號', bold=True, size=11)
cell_text(tobj.rows[0].cells[1], '單元學習目標', bold=True, size=11)
cell_text(tobj.rows[0].cells[2], '調整後學習目標', bold=True, size=11)
for c in range(3): set_cell_bg(tobj.rows[0].cells[c], 'D9E1F2')

goals = [
    ('1', '能正確辨識台灣常用幣值（1、5、10、50、100元）',
     '1-1 能指出10元並說出面額\n1-2 能區分硬幣與鈔票'),
    ('2', '能計算指定金額的硬幣/鈔票組合',
     '2-1 能數出5枚10元硬幣共50元\n2-2 在計算機輔助下完成加總'),
    ('3', '能從錢包中選取正確金額完成付款',
     '3-1 在語音提示下選取一枚指定硬幣\n3-2 在輔助點擊模式下完成付款流程'),
    ('4', '能在自動販賣機/超市模擬中完成購物全流程',
     '4-1 在簡單模式下完成自動販賣機購物\n4-2 能說出流程步驟（至少3步）'),
    ('5', '能計算付款後應找回的零錢',
     '5-1 能在提示下確認找零金額正確與否'),
    ('6', '能類化技能於真實消費情境',
     '6-1 在教師陪同下於便利商店完成購物'),
    ('7', '能辨識超市中不同商品的標價與自身預算的關係',
     '7-1 能判斷現有金額能否購買指定商品'),
]
for i, (code, goal, adj_goal) in enumerate(goals):
    cell_text(tobj.rows[i+1].cells[0], code, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(tobj.rows[i+1].cells[1], goal, size=11)
    cell_multiline(tobj.rows[i+1].cells[2], adj_goal.split('\n'), size=11)
    tobj.rows[i+1].cells[0].width = Cm(1.0)
    tobj.rows[i+1].cells[1].width = Cm(7.0)
    tobj.rows[i+1].cells[2].width = Cm(7.5)

doc.add_paragraph()

# ── 教學活動設計
add_heading(doc, '教學活動設計', 2, 12, (0,51,102))

act_data = [
    # (目標代號, 教學活動, 時間, 課程調整)
    [
        ('節次', '第一節：認識台灣錢幣（C1、C2 單元）', '45分鐘', ''),
        ('1、1-1\n1-2', '【引起動機】\n展示真實硬幣及鈔票，讓學生觸摸觀察。\n問：「你有用過這個嗎？在哪裡？」\n播放日常消費影片片段（超市、便利商店）。', '8分', '環境：傳遞實體錢幣\n評量：觀察、口頭回應'),
        ('1、1-1\n1-2', '【發展活動一：辨識錢幣 C1 簡單模式】\n開啟金錢小達人 C1「認識錢幣」（簡單模式）。\n教師投影示範操作，說明每種幣的特徵。\n學生輪流操作：看圖點選對應面額，語音同步。\n困難模式：聽聲音選出正確錢幣（甲乙丙分層）。', '15分', '歷程：多層次教學\n甲丁：簡單+輔助點擊\n乙：普通模式\n丙：困難模式\n評量：口頭評量'),
        ('2、2-1\n2-2', '【發展活動二：數錢練習 C2 單元】\n開啟 C2「數錢」，學生逐枚點擊螢幕上的硬幣，\n語音同步播報「一元、二元…」。\n教師依據學生表現適時停頓，提問確認理解。', '12分', '內容：甲丁限1、5、10元\n乙丙：全面額\n評量：實作評量'),
        ('1、2', '【統整活動】\n教師出示真實錢幣，學生說出面額及個數。\n發放 C1 作業單（認識錢幣）：圈選正確面額。\n總結：分享今天認識了哪些錢幣。', '10分', '評量：紙本評量（圈選式）\n甲丁：3選1\n乙丙：6選1'),
    ],
    [
        ('節次', '第二節：付款技能練習（C4、C5 單元）', '45分鐘', ''),
        ('3、3-1', '【引起動機】\n播放「在超市結帳」短片，觀察收銀台付款過程。\n問：「如果你要買一瓶 18 元的水，你會拿哪幾枚硬幣？」', '8分', '評量：口頭評量'),
        ('3、3-1\n5、5-1', '【發展活動一：付款練習 C4 單元】\n開啟 C4「付款」（普通模式），學生從錢包\n拖曳正確金額至付款區，完成付款。\n教師逐題確認，對錯誤次數多的學生切換為簡單模式。\n特色操作：嘗試「超額付款」觀察找零流程。', '15分', '歷程：提示線索（3次後）\n甲丁：簡單模式\n乙：普通模式\n丙：困難模式\n評量：實作評量'),
        ('5、5-1', '【發展活動二：夠不夠 C5 單元】\n開啟 C5「夠不夠」，學生判斷錢包金額\n能否買下商品，需說出理由。\n教師引導討論：「如果不夠，你會怎麼做？」', '12分', '內容：金額限50元以內（甲丁）\n評量：口頭 + 實作'),
        ('3、4', '【統整活動】\n模擬情境：教師展示商品標價卡，\n學生從準備好的硬幣組合中選取付款。\n發放 C4 作業單。', '10分', '評量：實作 + 紙本\n甲丁：老師提供正確組合選項'),
    ],
    [
        ('節次', '第三節：自動販賣機模擬操作（A1 單元）', '45分鐘', ''),
        ('4、4-1', '【引起動機】\n播放真實自動販賣機操作影片。\n問：「你有自己買過飲料嗎？怎麼投幣的？」\n預習流程：師生共同說出5步驟流程。', '8分', '評量：口頭評量'),
        ('4、4-1', '【發展活動一：教師示範 A1 完整流程】\n教師投影示範：選飲料→投幣→確認購買→\n取飲料→取找零的完整流程。\n說明各步驟的提示（黃色光暈、語音）。', '7分', '環境：大螢幕投影\n歷程：工作分析'),
        ('4、4-1\n5、5-1', '【發展活動二：學生輪流操作 A1】\n學生輪流操作，教師個別指導（每人2-3輪）。\n甲丁生：簡單模式+輔助點擊；乙生：普通模式；\n丙生：困難模式，嘗試自行計算找零。\n完成後播放動畫與語音鼓勵，學生看到煙火效果。', '20分', '歷程：多層次教學\n甲丁：輔助點擊\n評量：實作評量\n觀察：完成步驟數'),
        ('4', '【統整活動】\n學生分享操作心得：「最難的步驟是？」\n教師給予個別口頭回饋，貼上完成貼紙（獎勵）。\n預告：下節課要去「超市」購物！', '10分', '評量：口頭評量\n環境：獎勵貼紙'),
    ],
    [
        ('節次', '第四節：超市購物模擬與類化（A4 單元）', '45分鐘', ''),
        ('6、7', '【引起動機】\n展示超市環境照片，討論「便利商店、超市、菜市場有什麼不同？」\n說明本節目標：用有限的「零用錢」在超市選購。', '8分', '評量：觀察'),
        ('6、7\n4-1', '【發展活動一：A4 超市購物情境練習】\n開啟 A4「超市購物」，設定便利商店情境。\n學生進行「指定任務」（教師指定商品，學生完成購買）。\n逐步提升難度：增加商品種類、提高金額。', '17分', '內容：乙丙可嘗試組合購買\n甲丁：限單一商品\n評量：實作評量'),
        ('7', '【發展活動二：「我的100元怎麼花？」模擬】\n教師給每位學生「100元零用錢」（虛擬）。\n學生在 A4 自選模式下挑選商品，\n觀察「總金額不能超過100元」的限制。\n討論：「如果超過了怎麼辦？要放回哪樣商品？」', '10分', '歷程：合作學習（兩人討論）\n評量：口頭評量\n甲丁：教師共同操作'),
        ('4、6', '【統整與類化】\n討論：今天學到的技能，在真實生活中什麼時候會用到？\n佈置作業：「下次和家人去超市，自己選一樣商品付款」。\n發放 A4 作業單、自評表。', '10分', '評量：紙本 + 自評表\n家庭連結：家長共同練習任務單'),
    ],
]

for lesson in act_data:
    t_act = doc.add_table(rows=len(lesson)+1, cols=4)
    t_act.style = 'Table Grid'
    # 標頭列
    headers = ['目標代號', '教學活動', '時間', '課程調整（內容/歷程/環境/評量）']
    for ci, h in enumerate(headers):
        cell_text(t_act.rows[0].cells[ci], h, bold=True, size=10)
        set_cell_bg(t_act.rows[0].cells[ci], 'D9E1F2')
    t_act.rows[0].cells[0].width = Cm(1.5)
    t_act.rows[0].cells[1].width = Cm(8.5)
    t_act.rows[0].cells[2].width = Cm(1.5)
    t_act.rows[0].cells[3].width = Cm(4.0)

    # 節次標題列（合併）
    node_row = lesson[0]
    nr = t_act.rows[1]
    nr.cells[0].merge(nr.cells[3])
    cell_text(nr.cells[0], node_row[1], bold=True, size=11)
    set_cell_bg(nr.cells[0], 'EBF4FF')

    # 活動列
    for ri, (code, activity, time, adj) in enumerate(lesson[1:], 2):
        row = t_act.rows[ri]
        cell_text(row.cells[0], code, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_multiline(row.cells[1], activity.split('\n'), size=10)
        cell_text(row.cells[2], time, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_multiline(row.cells[3], adj.split('\n'), size=10)
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(8.5)
        row.cells[2].width = Cm(1.5)
        row.cells[3].width = Cm(4.0)

    doc.add_paragraph()

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ═══════════════════════════════════════════════════════
# 附表三之一 ─ 教學活動設計（單元二）：ATM
# ═══════════════════════════════════════════════════════

add_heading(doc, '附表三之一　特殊教育教學活動設計', 1, 14)
add_heading(doc, '教學設計（二）：ATM 操作與日常金融技能', 2, 13, (0,102,51))
add_para(doc, '——以 ATM 提款與轉帳為情境（A5 單元）', size=12)
doc.add_paragraph()

# 基本資訊表
ti2 = doc.add_table(rows=8, cols=4)
ti2.style = 'Table Grid'

def ti2_row(row_idx, c0, c1, c2=None, c3=None, merge12=False, merge_all=False):
    row = ti2.rows[row_idx]
    cell_text(row.cells[0], c0, bold=True, size=11)
    set_cell_bg(row.cells[0], 'D9E1F2')
    row.cells[0].width = Cm(2.8)
    if merge_all:
        row.cells[1].merge(row.cells[3])
        cell_multiline(row.cells[1], c1 if isinstance(c1, list) else [c1], size=11)
    elif merge12:
        row.cells[1].merge(row.cells[2])
        cell_text(row.cells[1], c1, size=11)
        cell_text(row.cells[3], c3 or '', size=11)
    elif c2 is not None:
        cell_text(row.cells[1], c1, size=11)
        cell_text(row.cells[2], c2, bold=True, size=11)
        set_cell_bg(row.cells[2], 'D9E1F2')
        cell_text(row.cells[3], c3 or '', size=11)
    else:
        row.cells[1].merge(row.cells[3])
        cell_text(row.cells[1], c1, size=11)
    for i in range(4): row.cells[i].width = Cm(2.8) if i in [0,2] else Cm(5.2)

ti2_row(0, '領域/科目', '■ 數學領域　■ 融入特殊需求領域（生活管理）', merge_all=False)
ti2_row(1, '形式', '■ 融入　生活技能/轉銜教育　領域/科目', merge_all=False)
ti2_row(2, '實施型態', [
    '■ 身心障礙分散式資源班　■ 集中式特教班',
], merge_all=True)
ti2_row(3, '單元名稱', '自立生活金融技能——ATM 操作實戰',
        '版本', '自編（金錢小達人 A5 單元）')
ti2_row(4, '學習階段/年級', '第四至第五階段 / 八至十二年級',
        '教學設計者', '特教教師（請填入姓名）')
ti2_row(5, '學習功能', '■ 輕微缺損',
        '教學時間', '共 3 節 / 每節 45 分鐘')
ti2_row(6, '特殊需求', '■ 生活管理　■ 輔助科技應用　■ 職業教育　■ 社會技巧', merge_all=True)
ti2_row(7, '議題融入', '■ 科技教育　■ 資訊教育　■ 生涯規劃教育　■ 安全教育', merge_all=True)

doc.add_paragraph()

# 設計理念
add_heading(doc, '設計理念', 2, 12, (0,51,102))
add_para(doc, '對於特殊需求青少年而言，ATM 操作是邁向獨立生活的關鍵里程碑之一。然而，面對真實 ATM 的多步驟流程、複雜介面及外在環境壓力，許多學生在未充分練習前便產生明顯焦慮與逃避行為。本教學設計利用金錢小達人 A5「ATM 模擬器」，提供低壓力、可反覆練習的數位環境，讓學生在正式接觸真實 ATM 前建立充分的程序性記憶與操作信心，進而達成「數位模擬→角色扮演→社區實地」的漸進式轉銜學習歷程，符合轉銜教育及功能性技能訓練的核心精神。', size=12)
doc.add_paragraph()

# 學生能力描述（ATM 版）
add_heading(doc, '學生能力描述', 2, 12, (0,51,102))
ts2 = doc.add_table(rows=4, cols=5)
ts2.style = 'Table Grid'
for i, h in enumerate(['學生', '特教類別', '一般現況能力', 'ATM相關能力現況', '課程調整']):
    cell_text(ts2.rows[0].cells[i], h, bold=True, size=10)
    set_cell_bg(ts2.rows[0].cells[i], 'D9E1F2')

students2 = [
    ('甲生', '輕度智能障礙',
     '數字識讀能力正常；操作電腦有基礎；對新環境適應需要時間',
     '知道 ATM 可以領錢，但從未獨立操作過；對「輸入密碼」步驟感到緊張',
     '簡化：簡單模式；分解步驟逐一練習；密碼輸入使用固定4碼練習'),
    ('乙生', '自閉症（輕度）',
     '認知及數學能力接近同齡；對電子設備高度熟悉；固著於已知程序',
     '曾看過父母使用 ATM；能說出部分步驟，但遇到「選擇選單」時易混淆',
     '歷程調整：提示流程卡；普通模式；針對選單頁面進行額外說明'),
    ('丙生', '輕度智能障礙',
     '動作協調良好；語言能力中等；對金錢使用有高度學習動機',
     '能完成付款，但對 ATM 功能（提款/存款/查詢）仍不熟悉',
     '普通模式；目標：完整完成提款及查詢兩種功能'),
]
for r, (s, cat, gen, domain, adj) in enumerate(students2):
    for c, t in enumerate([s, cat, gen, domain, adj]):
        cell_text(ts2.rows[r+1].cells[c], t, size=10)

doc.add_paragraph()

# 學習目標（ATM版）
add_heading(doc, '學習目標', 2, 12, (0,51,102))
tobj2 = doc.add_table(rows=7, cols=3)
tobj2.style = 'Table Grid'
for i, h in enumerate(['代號', '單元學習目標', '調整後學習目標']):
    cell_text(tobj2.rows[0].cells[i], h, bold=True, size=11)
    set_cell_bg(tobj2.rows[0].cells[i], 'D9E1F2')
goals2 = [
    ('1', '能說出 ATM 操作的基本步驟（插卡→密碼→選功能→操作→取卡）',
     '1-1 能配合流程卡說出至少3個步驟'),
    ('2', '能在數位模擬器中獨立完成 ATM 提款操作',
     '2-1 在簡單模式+輔助點擊下完成一次提款流程'),
    ('3', '能完成存款與查詢餘額操作',
     '3-1 在提示下選擇正確的交易選項'),
    ('4', '能在教師陪同下於真實 ATM 完成提款（轉銜目標）',
     '4-1 配合流程卡，在教師示範後嘗試一次操作'),
    ('5', '能識別 ATM 使用的安全注意事項',
     '5-1 能說出至少2項安全注意事項（不讓陌生人看密碼等）'),
    ('6', '能判斷存款、提款、轉帳、查詢四種功能的適用情境',
     '6-1 能在情境題中選出正確的 ATM 操作類型'),
]
for i, (code, goal, adj) in enumerate(goals2):
    cell_text(tobj2.rows[i+1].cells[0], code, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(tobj2.rows[i+1].cells[1], goal, size=11)
    cell_text(tobj2.rows[i+1].cells[2], adj, size=11)
    tobj2.rows[i+1].cells[0].width = Cm(1.0)
    tobj2.rows[i+1].cells[1].width = Cm(7.0)
    tobj2.rows[i+1].cells[2].width = Cm(7.5)

doc.add_paragraph()

# 教學活動設計（ATM 3節）
add_heading(doc, '教學活動設計', 2, 12, (0,51,102))

atm_lessons = [
    [
        ('節次', '第一節：ATM 介紹與提款模擬練習', '45分鐘', ''),
        ('1', '【引起動機】\n展示 ATM 照片（街景、銀行），問：「你知道這是什麼嗎？\n你有看過家人使用嗎？用來做什麼？」\n展示「ATM 操作步驟流程卡」（圖文並茂），說明本節目標。', '10分', '環境：流程卡\n評量：口頭評量'),
        ('1、2', '【發展活動一：教師示範 A5 完整提款流程】\n投影開啟金錢小達人 A5 簡單模式，教師完整示範：\n插卡→輸入PIN→選「提款」→選金額→取現金→\n選「明細表」→取明細→退卡。\n每個步驟配合流程卡說明，學生跟讀步驟名稱。', '10分', '歷程：工作分析\n環境：大螢幕投影\n評量：觀察'),
        ('2、2-1', '【發展活動二：學生輪流操作（提款）】\n甲生：簡單模式+輔助點擊（每次點擊自動執行下一步）。\n乙丙生：普通模式，嘗試自行完成提款流程。\n每位學生完成2輪，教師個別巡視指導。\n完成後軟體顯示成績統計與煙火動畫，給予正向回饋。', '18分', '歷程：多層次教學\n甲：輔助點擊\n評量：實作評量\n記錄：觀察記錄表'),
        ('1', '【統整活動】\n配合流程卡，學生輪流說出提款步驟。\n發放 A5 作業單（ATM 流程排序題）。\n預告：下節練習「存款」與「查詢」。', '7分', '評量：口頭+紙本\n甲：圖片排序'),
    ],
    [
        ('節次', '第二節：存款、查詢與轉帳模擬練習', '45分鐘', ''),
        ('1、3', '【引起動機】\n情境題：「你媽媽叫你去銀行存 500 元，你要選哪個功能？」\n複習流程卡，複習上節提款步驟。', '8分', '評量：口頭評量'),
        ('3、3-1\n6、6-1', '【發展活動一：存款模擬 A5】\n開啟 A5 簡單/普通模式，選擇「存款」功能，\n學生練習點選鈔票面額（現金出口操作）→確認→退卡。\n討論：「存款和提款有什麼不一樣？」', '15分', '歷程：提示線索\n評量：實作評量\n甲：教師逐步引導'),
        ('6、6-1', '【發展活動二：查詢餘額 + 轉帳練習 A5】\n查詢餘額：學生操作「查詢」功能，觀察餘額顯示畫面。\n轉帳（乙丙生）：嘗試轉帳至指定帳號，\n注意：輸入行碼→帳號→金額→確認流程。\n（甲生：僅完成查詢，不要求轉帳）', '12分', '內容：甲不要求轉帳\n評量：實作+觀察\n丙：困難模式嘗試'),
        ('5、5-1\n6', '【統整活動：ATM 安全我最行】\n討論：使用 ATM 要注意哪些安全事項？\n「密碼要保密」「不讓陌生人看」「附近沒人再操作」等。\n製作「ATM 安全守則」小海報（小組合作）。', '10分', '歷程：合作學習\n評量：實作評量'),
    ],
    [
        ('節次', '第三節：角色扮演與轉銜準備', '45分鐘', ''),
        ('1～5', '【引起動機】\n播放真實 ATM 環境影片（銀行門口街景），\n討論：「如果你一個人去，你有信心嗎？有什麼擔心？」\n說明本節目標：模擬獨立使用 ATM 的情境。', '8分', '評量：觀察\n環境：移除教師支持漸次獨立'),
        ('2、3\n4', '【發展活動一：A5 完整流程個人挑戰】\n學生各自選擇一種交易類型（提款/存款/查詢/轉帳），\n嘗試在最少提示下獨立完成全流程。\n教師僅在確實卡住時給予提示，其餘旁觀記錄。', '15分', '評量：實作評量（記錄提示次數）\n甲：仍可使用輔助點擊\n乙丙：盡量不依賴提示'),
        ('4、4-1', '【發展活動二：角色扮演 ATM 體驗】\n教室佈置模擬「銀行街」情境（帶 ATM 機台照片海報）。\n兩人一組：一人扮演「使用者」、一人扮演「引導夥伴」。\n使用者操作 A5 軟體，引導夥伴依照流程卡給予口頭提示。\n角色互換，確保每人均有完整操作體驗。', '15分', '歷程：合作學習+社會技巧\n評量：觀察+同儕評量\n環境：模擬銀行情境'),
        ('4、5', '【統整活動：轉銜計畫討論】\n教師說明下週將安排「社區 ATM 參訪體驗」。\n發放「ATM 操作練習護照」（家長版流程卡），\n請家長在家陪同練習並記錄。\n完成獎勵系統點數累積，頒發「金融小達人」貼紙。', '7分', '評量：口頭+自評表\n家庭連結：護照任務單\n獎勵：點數+貼紙'),
    ],
]

for lesson in atm_lessons:
    t_atm = doc.add_table(rows=len(lesson)+1, cols=4)
    t_atm.style = 'Table Grid'
    headers = ['目標代號', '教學活動', '時間', '課程調整（內容/歷程/環境/評量）']
    for ci, h in enumerate(headers):
        cell_text(t_atm.rows[0].cells[ci], h, bold=True, size=10)
        set_cell_bg(t_atm.rows[0].cells[ci], 'D9E1F2')
    t_atm.rows[0].cells[0].width = Cm(1.5)
    t_atm.rows[0].cells[1].width = Cm(8.5)
    t_atm.rows[0].cells[2].width = Cm(1.5)
    t_atm.rows[0].cells[3].width = Cm(4.0)

    # 節次行
    nr = t_atm.rows[1]
    nr.cells[0].merge(nr.cells[3])
    cell_text(nr.cells[0], lesson[0][1], bold=True, size=11)
    set_cell_bg(nr.cells[0], 'EBF4FF')

    for ri, (code, activity, time, adj) in enumerate(lesson[1:], 2):
        row = t_atm.rows[ri]
        cell_text(row.cells[0], code, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_multiline(row.cells[1], activity.split('\n'), size=10)
        cell_text(row.cells[2], time, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_multiline(row.cells[3], adj.split('\n'), size=10)
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(8.5)
        row.cells[2].width = Cm(1.5)
        row.cells[3].width = Cm(4.0)

    doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 備注 + 結語
# ═══════════════════════════════════════════════════════

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
add_heading(doc, '附件說明', 1, 13)
add_para(doc, '本文件包含以下競賽申請資料：', size=12, bold=True)
items = [
    '附表一：競賽申請表（請手填姓名、身分證字號、服務單位、聯絡資訊並簽名）',
    '附表三：創作說明（本文件第二節）',
    '附表三之一（一）：特殊教育教學活動設計——貨幣辨識與日常消費技能（4節）',
    '附表三之一（二）：特殊教育教學活動設計——ATM 操作與日常金融技能（3節）',
    '（附表二著作權聲明及授權書、附表四共同作者同意書請另行填寫）',
]
for item in items:
    p = add_para(doc, f'　{item}', size=11)

doc.add_paragraph()
add_heading(doc, '注意事項', 1, 13)
notes = [
    '1. 本軟體為網頁應用程式（HTML5/CSS3/JavaScript），無需安裝，開啟 Chrome 或 Edge 瀏覽器即可執行。',
    '2. 適用裝置：Windows 電腦（推薦）、iPad、Android 平板，螢幕尺寸建議 9.7 吋以上。',
    '3. 語音功能：需使用支援 Web Speech API 的瀏覽器；推薦語音：Microsoft Yating（雅婷）或 Google 國語。',
    '4. 作業單：各單元均有對應作業單，可依學生程度調整題型與題數後列印使用。',
    '5. 獎勵系統：支援教師上傳學生照片、累積點數、設定獎項，可搭配 IEP 正增強策略使用。',
    '6. 軟體原始碼及完整使用說明詳見隨附光碟。',
]
for note in notes:
    add_para(doc, note, size=11)

# ─── 儲存 ────────────────────────────────────────────────────

import os
out_path = 'C:/Users/super/Downloads/money_tutor0311/money_tutor/金錢小達人_第八屆特教教材競賽申請文件.docx'
doc.save(out_path)
size = os.path.getsize(out_path)
print(f'✅ 文件已儲存：{out_path}')
print(f'   檔案大小：{size:,} bytes ({size//1024} KB)')
