# -*- coding: utf-8 -*-
"""
建立金錢小達人競賽申請完整文件 (Word)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "C:/Users/super/Downloads/money_tutor0311/money_tutor/金錢小達人_競賽申請完整文件.docx"

# ──────────────────────────────────────────
# 輔助函式
# ──────────────────────────────────────────

def set_font(run, size=12, bold=False, zh_font="標楷體", en_font="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font)

def set_para_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=6,
                    line_spacing=1.5):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing

def add_heading(doc, text, level=1, size=14):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=size, bold=True)
    set_para_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
    return para

def add_normal_para(doc, text, size=12, bold=False, indent=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=size, bold=bold)
    set_para_format(para)
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.85)
    return para

def set_cell_font(cell, text=None, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """設定儲存格字型，若 text 為 None 則不重寫文字"""
    if text is not None:
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        set_font(run, size=size, bold=bold)
        set_para_format(para, align=align, space_before=2, space_after=2, line_spacing=1.3)
    else:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, size=size, bold=bold)

def set_cell_bg(cell, hex_color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_cell_paragraph(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """在 cell 中新增段落（不清空原有內容）"""
    para = cell.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=size, bold=bold)
    set_para_format(para, align=align, space_before=1, space_after=1, line_spacing=1.3)
    return para

def set_table_border(table):
    """設定表格所有框線"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_module.enum.text.WD_BREAK.PAGE)

# ──────────────────────────────────────────
# 主程式
# ──────────────────────────────────────────

import docx as docx_module

doc = Document()

# ── 頁面設定 ──
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(3.17)
section.bottom_margin = Cm(3.17)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ══════════════════════════════════════════
# 封面頁
# ══════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("教育部獎助研發特殊教育教材教具競賽")
set_font(run, size=16, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("競賽申請完整文件")
set_font(run, size=14, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("金錢小達人")
set_font(run, size=22, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("——互動式特殊教育生活數學教學軟體")
set_font(run, size=14, bold=False)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ("申請組別", "身心障礙類：國民小學組"),
    ("申請日期", "中華民國 115 年 3 月"),
]
for label, value in info_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f"【{label}】")
    set_font(run1, size=13, bold=True)
    run2 = p.add_run(f"　{value}")
    set_font(run2, size=13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 分頁
doc.add_page_break()

# ══════════════════════════════════════════
# 第一部分：附表三 — 創作說明
# ══════════════════════════════════════════

add_heading(doc, "附表三　創作說明", level=1, size=14)

# 建立表格
t1 = doc.add_table(rows=11, cols=2)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(t1)

# 設定欄寬（左欄 3.5cm，右欄約 13cm）
LEFT_W = 3.5
RIGHT_W = 13.0

rows_data = [
    ("參賽組別",   "■ 身心障礙類：■ 國民小學組"),
    ("作品名稱",   "金錢小達人——互動式特殊教育生活數學教學軟體"),
    ("適用對象",   "國民小學一至六年級身心障礙類學生（含輕度至中度智能障礙、自閉症、學習障礙等特殊需求學生）"),
    ("作品用途",   "本軟體為一套以網頁為基礎的互動式特殊教育生活數學電腦輔助教學軟體，提供教師、家長及特殊需求學生使用，透過 18 個生活情境單元及 6 個數學基礎單元，系統性培養學生的貨幣辨識、金額計算、找零付款及日常消費等生活數學能力。"),
    ("設計動機",   "特殊需求學生在「生活數學」領域的學習，往往面臨抽象概念難以內化、缺乏真實情境練習機會、個別差異大且坊間缺乏符合課綱的系統性教材等困境。研究者長期觀察教學現場，發現學生對於「金錢」的實際操作——如投幣、找零、ATM 操作等——雖在日常生活中不可或缺，卻因練習機會不足而缺乏信心與能力。基於此，研究者自主研發本套軟體，結合課程通用設計原則（Universal Design for Learning, UDL），以「先備知識建立→基礎認知→生活應用」的學習序列，設計一套適用於不同程度特殊需求學生的互動式教學系統。"),
    ("製作費用",   "軟體開發費：新臺幣 0 元（研究者自行開發）；硬體設備（平板／電腦）：各校自行準備"),
    ("經費來源",   "自籌"),
]

# 多行文字的特殊列
content_overview = [
    "本軟體共設計 18 個完整教學單元，分為三大系列：",
    "",
    "【F 系列：數學基礎】（6 個單元）",
    "F1 一對一對應、F2 唱數、F3 數字認讀、F4 數字排序、F5 量比較、F6 數的組成。以豐富的視覺化圖示與多感官互動，建立學生的數概念基礎。",
    "",
    "【C 系列：貨幣認知】（6 個單元）",
    "C1 認識錢幣、C2 數錢、C3 換錢、C4 付款、C5 夠不夠、C6 找零。從辨識台灣常見硬幣與鈔票、數算金額、等值兌換，到付款與找零，逐步建立完整的貨幣認知體系。",
    "",
    "【A 系列：生活應用】（6 個單元）",
    "A1 自動販賣機、A2 理髮廳自助機、A3 麥當勞點餐機、A4 超市購物、A5 ATM 操作、A6 火車售票機。完整模擬六種真實生活消費場景，讓學生在安全的數位環境中反覆練習。",
    "",
    "【製作過程】",
    "本軟體以 HTML5 + CSS3 + JavaScript 開發，採用無框架純前端架構，支援電腦、平板及手機多裝置使用。所有教學素材（圖片、音效、語音）均為原創或合法授權使用。開發歷程約 3 個月，共計 18 個主程式檔案（各約 3,000～8,000 行），及 30 個以上輔助模組。",
]

usage_effect = [
    "【使用方式】",
    "• 教師開啟瀏覽器即可使用，無需安裝軟體",
    "• 三種難度（簡單／普通／困難）與兩種以上任務類型，教師可依學生程度彈性選擇",
    "• 「輔助點擊模式」：畫面任何位置點一下即執行下一步，適用於動作協調困難或認知負荷較高的學生",
    "• 語音系統全程提供中文語音引導，支援 Google 語音與 Windows 語音",
    "• 作業單系統：可一鍵產生對應的印刷式紙本練習卷，提供 10 種以上題型",
    "",
    "【教學成效】",
    "根據試用回饋：",
    "• 學生對「真實情境模擬」高度投入，尤其 A1（販賣機）、A5（ATM）反應熱烈",
    "• 輔助點擊模式有效降低動作困難學生的操作門檻，達成「我也可以獨立完成」的成就感",
    "• 教師反映「作業單與教學軟體完全對應」節省大量自製教材時間",
    "• 三難度設計使同一班級不同程度學生均能找到適合的挑戰",
]

appendix_text = "見後附特殊教育教學活動設計（附表三之一）及軟體操作截圖"

# 填入各列
def fill_row(row, label, value_lines=None, value_str=None, label_bg="D9E1F2"):
    lc = row.cells[0]
    rc = row.cells[1]
    lc.width = Cm(LEFT_W)
    rc.width = Cm(RIGHT_W)
    set_cell_font(lc, label, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    lc.paragraphs[0].paragraph_format.space_before = Pt(3)
    lc.paragraphs[0].paragraph_format.space_after = Pt(3)
    set_cell_bg(lc, label_bg)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if value_str is not None:
        set_cell_font(rc, value_str, size=11)
    elif value_lines is not None:
        rc.text = ""
        for line in value_lines:
            p = rc.add_paragraph()
            run = p.add_run(line)
            is_header = line.startswith("【") and line.endswith("】")
            set_font(run, size=11, bold=is_header)
            set_para_format(p, space_before=1, space_after=1, line_spacing=1.3)
        # 移除空的第一段
        if rc.paragraphs and rc.paragraphs[0].text == "":
            para_elem = rc.paragraphs[0]._element
            para_elem.getparent().remove(para_elem)

# Row 0: 參賽組別
fill_row(t1.rows[0], "參賽組別", value_str="■ 身心障礙類：■ 國民小學組")
# Row 1: 作品名稱
fill_row(t1.rows[1], "作品名稱", value_str="金錢小達人——互動式特殊教育生活數學教學軟體")
# Row 2: 適用對象
fill_row(t1.rows[2], "適用對象", value_str="國民小學一至六年級身心障礙類學生（含輕度至中度智能障礙、自閉症、學習障礙等特殊需求學生）")
# Row 3: 作品用途
fill_row(t1.rows[3], "作品用途", value_str="本軟體為一套以網頁為基礎的互動式特殊教育生活數學電腦輔助教學軟體，提供教師、家長及特殊需求學生使用，透過 18 個生活情境單元及 6 個數學基礎單元，系統性培養學生的貨幣辨識、金額計算、找零付款及日常消費等生活數學能力。")
# Row 4: 設計動機
fill_row(t1.rows[4], "設計動機", value_str="特殊需求學生在「生活數學」領域的學習，往往面臨抽象概念難以內化、缺乏真實情境練習機會、個別差異大且坊間缺乏符合課綱的系統性教材等困境。研究者長期觀察教學現場，發現學生對於「金錢」的實際操作——如投幣、找零、ATM 操作等——雖在日常生活中不可或缺，卻因練習機會不足而缺乏信心與能力。基於此，研究者自主研發本套軟體，結合課程通用設計原則（Universal Design for Learning, UDL），以「先備知識建立→基礎認知→生活應用」的學習序列，設計一套適用於不同程度特殊需求學生的互動式教學系統。")
# Row 5: 內容概述
fill_row(t1.rows[5], "內容概述", value_lines=content_overview)
# Row 6: 使用說明及效果
fill_row(t1.rows[6], "使用說明\n及效果", value_lines=usage_effect)
# Row 7: 附錄
fill_row(t1.rows[7], "附錄", value_str=appendix_text)
# Row 8: 製作費用
fill_row(t1.rows[8], "製作費用", value_str="軟體開發費：新臺幣 0 元（研究者自行開發）；硬體設備（平板／電腦）：各校自行準備")
# Row 9: 經費來源
fill_row(t1.rows[9], "經費來源", value_str="自籌")
# Row 10: 備註空白列
fill_row(t1.rows[10], "備註", value_str="（無）")

doc.add_paragraph()

# ══════════════════════════════════════════
# 分頁 → 第二部分：附表三之一（單元一）
# ══════════════════════════════════════════
doc.add_page_break()

add_heading(doc, "附表三之一　特殊教育教學活動設計（單元一）", level=1, size=14)

# ── 基本資訊表格 ──
add_normal_para(doc, "壹、基本資訊", size=12, bold=True)

basic_info = [
    ("領域／科目", "■ 數學領域（融入特殊需求領域）"),
    ("形式", "■ 融入　生活技能　領域／科目"),
    ("實施型態", "■ 身心障礙分散式資源班　　■ 集中式特教班"),
    ("單元名稱", "認識錢幣與付款——自動販賣機模擬操作"),
    ("版本", "自編（金錢小達人軟體 A1 + C1～C4 單元）"),
    ("學習階段／年級", "第一至第三階段 ／ 一至六年級"),
    ("教學設計者", "（請填入教師姓名）"),
    ("學習功能", "■ 輕微缺損　　■ 嚴重缺損"),
    ("教學時間", "共 3 節 ／ 每節 40 分鐘"),
    ("教學地點", "電腦教室或設有平板電腦之教室"),
    ("特殊需求", "■ 生活管理　　■ 輔助科技應用"),
    ("議題融入", "■ 科技教育　　■ 資訊教育　　■ 生涯規劃教育"),
]

t_basic = doc.add_table(rows=len(basic_info), cols=2)
set_table_border(t_basic)
t_basic.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, val) in enumerate(basic_info):
    fill_row(t_basic.rows[i], label, value_str=val)

doc.add_paragraph()

# ── 設計理念 ──
add_normal_para(doc, "貳、設計理念", size=12, bold=True)
design_concept = (
    "本教學活動以「真實情境模擬→技能練習→實際應用」為設計架構，結合金錢小達人教學軟體之 A1 自動販賣機"
    "與 C1～C4 貨幣認知單元，讓身心障礙學生在安全的數位環境中，透過多感官互動（視覺動畫、聽覺語音、觸覺操作）"
    "反覆練習「辨識錢幣→數算金額→付款找零」的完整流程，達成功能性生活技能的學習目標。"
)
add_normal_para(doc, design_concept, size=11, indent=True)

doc.add_paragraph()

# ── 學生能力描述 ──
add_normal_para(doc, "參、學生能力描述（含三種程度）", size=12, bold=True)

stu_headers = ["學生", "特教類別", "一般現況能力", "領域能力表現", "課程調整"]
stu_rows = [
    ("甲生", "輕度智能障礙",
     "識字能力約小二程度，能辨別數字 1-100，可獨立完成日常生活作息",
     "能辨識 1、5、10 元硬幣，但無法正確計算組合金額",
     "簡化：使用簡單模式＋輔助點擊，限定硬幣種類"),
    ("乙生", "自閉症（輕度）",
     "認知能力接近同齡，但對陌生情境適應困難，固著於特定行為模式",
     "能數算硬幣，但遇到需找零時常拒絕嘗試",
     "減量：從熟悉的飲料品項開始，逐步擴展情境"),
    ("丙生", "學習障礙",
     "數學計算困難，但語言能力正常，學習動機良好",
     "知道硬幣面額但計算組合時常錯誤",
     "歷程調整：使用提示線索（黃色光暈提示）"),
]

t_stu = doc.add_table(rows=len(stu_rows)+1, cols=5)
set_table_border(t_stu)
t_stu.alignment = WD_TABLE_ALIGNMENT.CENTER

# 標題列
for j, h in enumerate(stu_headers):
    cell = t_stu.rows[0].cells[j]
    set_cell_font(cell, h, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(cell, "D9E1F2")

# 設定欄寬
col_widths_stu = [1.5, 2.5, 3.5, 3.5, 3.5]
for j, w in enumerate(col_widths_stu):
    for row in t_stu.rows:
        row.cells[j].width = Cm(w)

# 資料列
for i, (s, t, g, d, c) in enumerate(stu_rows):
    row = t_stu.rows[i+1]
    for j, text in enumerate([s, t, g, d, c]):
        set_cell_font(row.cells[j], text, size=10)

doc.add_paragraph()

# ── 設計依據 ──
add_normal_para(doc, "肆、設計依據", size=12, bold=True)
design_basis = [
    "• 核心素養：數-E-A2 具備基本的算術操作能力、理解數量關係，在生活情境中運用。",
    "• 學科領域學習重點：數-1-A-4、數-2-B-3（金錢計算相關）",
    "• 特殊需求領域：生活管理——金錢管理技能",
]
for line in design_basis:
    add_normal_para(doc, line, size=11)

doc.add_paragraph()

# ── 學習目標 ──
add_normal_para(doc, "伍、學習目標", size=12, bold=True)
add_normal_para(doc, "【單元學習目標】", size=11, bold=True)
unit_goals = [
    "1. 能正確辨識台灣常用錢幣（1元、5元、10元、50元、100元）",
    "2. 能從錢包中選取指定金額投入自動販賣機",
    "3. 能完成「選商品→投幣→確認→取物→取找零」完整購物流程",
    "4. 能在輔助下操作 ATM 進行簡單存提款操作",
]
for g in unit_goals:
    add_normal_para(doc, g, size=11)

add_normal_para(doc, "【調整後學習目標】", size=11, bold=True)
adj_goals = [
    "1-1 能指出 10 元硬幣並說出面額",
    "2-1 在語音提示下，從錢包選取一枚硬幣投入",
    "3-1 在輔助點擊模式下，完成購買一瓶飲料的完整流程",
]
for g in adj_goals:
    add_normal_para(doc, g, size=11)

doc.add_paragraph()

# ── 教學活動設計 ──
add_normal_para(doc, "陸、教學活動設計", size=12, bold=True)

# 第一節
add_normal_para(doc, "第一節（40 分鐘）：認識錢幣——C1 單元", size=12, bold=True)

lesson1_headers = ["目標代號", "教學活動", "時間", "課程調整"]
lesson1_rows = [
    ("1、1-1", "引起動機：展示真實錢幣，讓學生觀察、觸摸，問「你在哪裡看過這個？」", "5 分", "環境：傳遞實體錢幣"),
    ("1、1-1", "發展活動：開啟金錢小達人 C1「認識錢幣」，學生依序點擊各面額錢幣，聆聽語音介紹", "15 分", "歷程：輔助點擊模式；評量：口頭回應"),
    ("1、1-1", "發展活動：教師出示錢幣圖片，學生在軟體中找到對應面額並點擊", "10 分", "內容：簡化——僅限 1、5、10 元"),
    ("1", "統整活動：C1 作業單（認識錢幣）——圈選正確面額", "10 分", "評量：紙筆評量（調整為圈選式）"),
]

def make_lesson_table(doc, headers, rows, col_widths):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    set_table_border(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_cell_font(cell, h, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cell, "D9E1F2")
    for j, w in enumerate(col_widths):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    for i, row_data in enumerate(rows):
        for j, text in enumerate(row_data):
            set_cell_font(t.rows[i+1].cells[j], text, size=10)
    return t

make_lesson_table(doc, lesson1_headers, lesson1_rows, [1.5, 7.0, 1.5, 4.5])
doc.add_paragraph()

# 第二節
add_normal_para(doc, "第二節（40 分鐘）：數錢與付款——C2、C4 單元", size=12, bold=True)
lesson2_rows = [
    ("2、2-1", "複習：C1 複習問答，抽問個別學生", "5 分", "評量：口頭評量"),
    ("2、2-1", "開啟 C2「數錢」——簡單模式，學生逐一點擊螢幕上的硬幣進行數算", "15 分", "歷程：語音提示全程開啟"),
    ("2、3", "開啟 C4「付款」——學生在指定情境中選取正確金額完成付款", "15 分", "歷程：提示線索；內容：限定 50 元以內"),
    ("2、3", "統整：C4 作業單（付款練習）", "5 分", "評量：作業單"),
]
make_lesson_table(doc, lesson1_headers, lesson2_rows, [1.5, 7.0, 1.5, 4.5])
doc.add_paragraph()

# 第三節
add_normal_para(doc, "第三節（40 分鐘）：自動販賣機完整流程——A1 單元", size=12, bold=True)
lesson3_rows = [
    ("3、3-1", "引起動機：播放真實販賣機操作影片，討論「你有自己買過東西嗎？」", "5 分", "評量：觀察"),
    ("3、3-1", "開啟 A1「自動販賣機」簡單模式：教師示範一遍完整流程（選商品→投幣→確認→取物→取找零）", "10 分", "環境：大螢幕投影示範"),
    ("3、3-1", "學生輪流操作（輔助點擊模式或簡單模式）", "20 分", "歷程：輔助點擊；評量：實作評量"),
    ("3", "統整：分享今天學到什麼，填寫自我評量", "5 分", "評量：口頭＋自評"),
]
make_lesson_table(doc, lesson1_headers, lesson3_rows, [1.5, 7.0, 1.5, 4.5])

doc.add_paragraph()

# ══════════════════════════════════════════
# 分頁 → 第三部分：附表三之一（單元二）
# ══════════════════════════════════════════
doc.add_page_break()

add_heading(doc, "附表三之一　特殊教育教學活動設計（單元二）", level=1, size=14)

# ── 基本資訊 ──
add_normal_para(doc, "壹、基本資訊", size=12, bold=True)

basic_info2 = [
    ("領域／科目", "■ 數學領域"),
    ("形式", "■ 融入　生活技能　領域"),
    ("實施型態", "■ 集中式特教班"),
    ("單元名稱", "ATM 操作與日常金融技能"),
    ("版本", "自編（金錢小達人軟體 A5 單元）"),
    ("學習階段／年級", "第二至第三階段 ／ 三至六年級"),
    ("教學設計者", "（請填入教師姓名）"),
    ("學習功能", "■ 輕微缺損"),
    ("教學時間", "共 2 節 ／ 每節 40 分鐘"),
    ("教學地點", "電腦教室或設有平板電腦之教室"),
    ("特殊需求", "■ 生活管理　　■ 輔助科技應用"),
    ("議題融入", "■ 科技教育　　■ 資訊教育"),
]

t_basic2 = doc.add_table(rows=len(basic_info2), cols=2)
set_table_border(t_basic2)
t_basic2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, val) in enumerate(basic_info2):
    fill_row(t_basic2.rows[i], label, value_str=val)

doc.add_paragraph()

# ── 設計理念 ──
add_normal_para(doc, "貳、設計理念", size=12, bold=True)
concept2 = (
    "ATM 操作是現代社會重要的生活技能，特殊需求學生因缺乏練習機會，往往在真實情境中感到焦慮。"
    "本活動利用金錢小達人 A5 ATM 模擬器，提供低壓力、可反覆練習的數位環境，讓學生在正式操作真實 ATM 前"
    "建立足夠的程序性記憶與操作信心。本軟體支援提款、存款、查詢、轉帳四種情境，並設有三難度及輔助點擊模式，"
    "完整對應不同能力程度的學習需求。"
)
add_normal_para(doc, concept2, size=11, indent=True)

doc.add_paragraph()

# ── 學習目標 ──
add_normal_para(doc, "參、學習目標", size=12, bold=True)
goals2 = [
    "1. 能說出 ATM 操作的基本步驟（插卡→輸入密碼→選擇功能→取款／存款→取卡）",
    "2. 能在輔助下完成 ATM 簡單提款操作（軟體模擬）",
    "3. 能在教師陪同下完成真實 ATM 提款操作（實地練習）",
]
for g in goals2:
    add_normal_para(doc, g, size=11)

doc.add_paragraph()

# ── 教學活動設計 ──
add_normal_para(doc, "肆、教學活動設計", size=12, bold=True)

# 第一節
add_normal_para(doc, "第一節（40 分鐘）：ATM 介紹與模擬操作", size=12, bold=True)
lesson_atm1_headers = ["教學活動", "時間", "說明"]
lesson_atm1_rows = [
    ("引起動機：展示 ATM 圖片，討論使用經驗；提問「你在哪裡看過 ATM？」", "5 分", "觀察學生反應，了解先備經驗"),
    ("教師示範：以投影方式示範 A5 簡單模式完整提款流程（插卡→密碼→提款→取卡）", "10 分", "大螢幕示範；語音引導同步播放"),
    ("學生輪流操作（輔助點擊模式或簡單模式），教師逐一個別指導", "20 分", "記錄各生完成度；給予即時回饋"),
    ("統整：步驟卡製作——學生將操作步驟貼在正確順序", "5 分", "評量：步驟排序作業單"),
]
make_lesson_table(doc, lesson_atm1_headers, lesson_atm1_rows, [8.5, 1.5, 4.5])
doc.add_paragraph()

# 第二節
add_normal_para(doc, "第二節（40 分鐘）：四種功能練習與轉帳情境", size=12, bold=True)
lesson_atm2_rows = [
    ("複習：步驟卡問答，學生憑記憶說出各步驟", "5 分", "評量：口頭評量"),
    ("存款模擬練習：開啟 A5 存款模式，學生練習投入虛擬紙鈔並確認", "10 分", "歷程：語音提示；提示線索"),
    ("查詢餘額→轉帳模擬：依序練習查詢與轉帳兩種情境", "15 分", "輔助點擊模式；教師觀察記錄"),
    ("統整：討論真實 ATM 使用注意事項（安全、保密、環境選擇）", "10 分", "小組討論；製作安全守則海報"),
]
make_lesson_table(doc, lesson_atm1_headers, lesson_atm2_rows, [8.5, 1.5, 4.5])
doc.add_paragraph()

# ── 評量與回饋 ──
add_normal_para(doc, "伍、評量方式與課程調整策略", size=12, bold=True)

eval_headers = ["評量向度", "評量方式", "課程調整策略"]
eval_rows = [
    ("知識理解", "口頭問答、作業單（圈選、填空）", "減少題數、提供圖示選項"),
    ("技能操作", "實作評量（軟體操作完成度）", "允許多次嘗試、使用輔助點擊"),
    ("情意態度", "觀察紀錄、自我評量表", "調整評量情境、採用替代評量"),
]
make_lesson_table(doc, eval_headers, eval_rows, [3.0, 5.5, 6.0])

doc.add_paragraph()

# ── 備注 ──
add_normal_para(doc, "陸、備註", size=12, bold=True)
remarks = [
    "• 本教學活動設計可視學生實際能力彈性調整節次與難度。",
    "• 金錢小達人軟體各單元均提供對應之紙本作業單，教師可視需要印製使用。",
    "• 建議於課後提供家長居家練習指引，延伸學習成效至真實生活情境。",
    "• 如有學生需要更高輔助，可搭配觸控筆、大型觸控螢幕等輔助科技設備。",
]
for r in remarks:
    add_normal_para(doc, r, size=11)

# ── 儲存 ──
doc.save(OUTPUT_PATH)
print("File saved to: " + OUTPUT_PATH)
