# -*- coding: utf-8 -*-
"""
Markdown → DOCX 轉換器（金錢小達人競賽文件專用）
用法：python md_to_docx.py <input.md> <output.docx>
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 色彩常數 ────────────────────────────────────────────
C1   = (46, 125, 50)    # green-800  主色
C2   = (27, 94, 32)     # green-900
C4   = (1, 87, 155)     # blue-800   副色
C3   = (0, 77, 64)      # teal-900
C5   = (13, 71, 161)    # deep blue
GRAY = (90, 90, 90)
WHITE= (255,255,255)

H_FILL  = '2E7D32'      # 表頭背景（主綠）
H_FILL2 = '1565C0'      # 表頭背景（深藍）
BG1  = 'E8F5E9'
BG2  = 'E3F2FD'
BG3  = 'FFF8E1'

# ── 低階工具 ────────────────────────────────────────────
def set_font(run, size=12, bold=False, color=None, name='標楷體'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*(color or (0, 0, 0)))
    try:
        rPr = run._element.get_or_add_rPr()
        rPr.get_or_add_rFonts().set(qn('w:eastAsia'), name)
    except Exception:
        pass

def pf(para, before=0, after=4, line=1.5,
       align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0, first=0):
    f = para.paragraph_format
    f.space_before = Pt(before)
    f.space_after  = Pt(after)
    if line:
        f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        f.line_spacing = line
    f.alignment = align
    if indent: f.left_indent       = Cm(indent)
    if first:  f.first_line_indent = Cm(first)

def shade_tc(tc, fill):
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def shade_p(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def add_inline(para, text, size=12, bold=False, color=None, name='標楷體'):
    """支援 **bold** 語法的行內文字渲染"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            set_font(r, size, True, color, name)
        else:
            r = para.add_run(part)
            set_font(r, size, bold, color, name)

def div_line(doc, color='2E7D32', thick=6):
    p = doc.add_paragraph()
    pf(p, 2, 2, None)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(thick))
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)

def pgbr(doc):
    p = doc.add_paragraph()
    pf(p, 0, 0, None)
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)

# ── 中階工具 ────────────────────────────────────────────
def H(doc, text, level=1):
    sz = {1: 18, 2: 14, 3: 12, 4: 12, 5: 11}[min(level, 5)]
    co = {1: C1, 2: C4, 3: C3, 4: (0, 0, 100), 5: (80, 80, 80)}[min(level, 5)]
    al = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    bf = {1: 12, 2: 10, 3: 6, 4: 4, 5: 3}[min(level, 5)]
    p  = doc.add_paragraph()
    r  = p.add_run(text)
    set_font(r, sz, True, co)
    pf(p, bf, 4, 1.2, al)
    return p

def P(doc, text='', size=12, bold=False, color=None,
      before=0, after=4, line=1.5,
      align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0, bullet=None):
    para = doc.add_paragraph()
    ind  = indent
    if bullet:
        ind  = indent + 0.55
        text = bullet + '\u3000' + text
        para.paragraph_format.first_line_indent = Cm(-0.55)
    pf(para, before, after, line, align, ind)
    add_inline(para, text, size, bold, color)
    return para

def BOX(doc, text, fill=BG1, size=11, bold=False, before=4, after=6):
    para = doc.add_paragraph()
    pf(para, before, after, 1.4, WD_ALIGN_PARAGRAPH.JUSTIFY, 0.4)
    para.paragraph_format.right_indent = Cm(0.4)
    shade_p(para, fill)
    add_inline(para, text, size, bold)
    return para

def make_table(doc, headers, rows, cw=None, hfill=H_FILL, hs=10.5, ds=10.5):
    if not headers and not rows:
        return
    ncols = len(headers) if headers else len(rows[0])
    t = doc.add_table(1 + len(rows), ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade_tc(cell._tc, hfill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        pf(p, 2, 2, 1.2, WD_ALIGN_PARAGRAPH.CENTER)
        add_inline(p, h, hs, True, WHITE)
    # data rows
    for ri, row in enumerate(rows):
        fill = 'F1F8E9' if ri % 2 == 0 else 'FFFFFF'
        for ci, ct in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            shade_tc(cell._tc, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            pf(p, 2, 2, 1.3, WD_ALIGN_PARAGRAPH.LEFT)
            add_inline(p, str(ct), ds)
    if cw:
        for i, w in enumerate(cw):
            if i < ncols:
                for row in t.rows:
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ── Markdown 解析器 ────────────────────────────────────
def parse_table(lines):
    """解析 Markdown 表格，回傳 (headers, rows)"""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or re.match(r'^\|[-| :]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return [], []
    return rows[0], rows[1:]

def is_table_line(line):
    return line.strip().startswith('|')

def is_separator_line(line):
    return bool(re.match(r'^\|[-| :]+\|$', line.strip()))

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin  = Cm(2.54)
    sec.top_margin  = sec.bottom_margin = Cm(2.54)

    sty = doc.styles['Normal']
    sty.font.name = '標楷體'
    sty.font.size = Pt(12)
    try:
        sty.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    except Exception:
        pass

    i = 0
    n = len(lines)

    # 追蹤最近看到的 H1 (附表一、附表三 等)，
    # 遇到 H1「附表三」之後的 H1 插入分頁
    seen_h1_count = 0

    while i < n:
        raw = lines[i]
        line = raw.rstrip('\n')
        stripped = line.strip()

        # ── 空行 ──
        if not stripped:
            i += 1
            continue

        # ── 水平線 ---
        if re.match(r'^---+$', stripped) or re.match(r'^\*\*\*+$', stripped):
            div_line(doc)
            i += 1
            continue

        # ── 引用 > ──
        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            BOX(doc, text, fill=BG2, size=11)
            i += 1
            continue

        # ── 標題 # ──
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            # 移除標題中的 markdown 粗體
            title = re.sub(r'\*\*([^*]+)\*\*', r'\1', title)
            if level == 1:
                if seen_h1_count > 0:
                    pgbr(doc)
                H(doc, title, 1)
                div_line(doc)
                seen_h1_count += 1
            elif level == 2:
                H(doc, title, 2)
            elif level == 3:
                H(doc, title, 3)
            elif level == 4:
                H(doc, title, 4)
            else:
                H(doc, title, 5)
            i += 1
            continue

        # ── 表格 ──
        if is_table_line(stripped):
            table_lines = []
            while i < n and is_table_line(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_table(table_lines)
            if headers:
                # 根據欄數估算欄寬
                ncols = len(headers)
                avail = 15.5  # 可用寬度 (cm)
                if ncols == 2:
                    cw = [4, avail - 4]
                elif ncols == 3:
                    cw = [avail / 3] * 3
                elif ncols == 4:
                    cw = [2, avail - 6, 2, 2]
                elif ncols == 5:
                    cw = [1.5, 2.5, 3, 6, 2.5]
                elif ncols == 6:
                    cw = [1.5, 2.3, 3, 5.5, 1.7, 1.5]
                else:
                    cw = [avail / ncols] * ncols
                make_table(doc, headers, rows, cw=cw, hfill=H_FILL2)
            continue

        # ── 無序列表 - 或 * ──
        if re.match(r'^[-*]\s', stripped):
            items = []
            while i < n and re.match(r'^[\s]*[-*]\s', lines[i]):
                items.append(re.sub(r'^[\s]*[-*]\s', '', lines[i]).strip())
                i += 1
            for item in items:
                P(doc, item, indent=0.3, before=1, after=2, line=1.4, bullet='●')
            continue

        # ── 有序列表 1. ──
        if re.match(r'^\d+\.\s', stripped):
            items = []
            while i < n and re.match(r'^\d+\.\s', lines[i].strip()):
                m2 = re.match(r'^\d+\.\s+(.*)', lines[i].strip())
                if m2:
                    items.append(m2.group(1).strip())
                i += 1
            for j, item in enumerate(items, 1):
                P(doc, f'{j}. {item}', indent=0.3, before=1, after=2, line=1.4)
            continue

        # ── 普通段落 ──
        # 收集連續普通行（直到空行/標題/列表/表格/引用/水平線）
        para_lines = []
        while i < n:
            l = lines[i].rstrip('\n').strip()
            if (not l or
                l.startswith('#') or
                l.startswith('|') or
                l.startswith('>') or
                re.match(r'^---+$', l) or
                re.match(r'^[-*]\s', l) or
                re.match(r'^\d+\.\s', l)):
                break
            para_lines.append(l)
            i += 1
        text = ' '.join(para_lines).strip()
        if text:
            P(doc, text, before=0, after=4, line=1.5)

    doc.save(docx_path)
    print(f'OK: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('用法：python md_to_docx.py <input.md> <output.docx>')
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
